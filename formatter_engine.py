"""
Motor de formatação — Revista Educar Mais
Lê um artigo .docx enviado por um autor e devolve um .docx já formatado
seguindo o layout real do Modelo.docx da revista:

- Cabeçalho em duas colunas: logo + título (PT/EN/ES) à esquerda,
  autores à direita, cada um com ícone do ORCID linkado e nota de
  rodapé com a biografia.
- Corpo do texto (RESUMO em diante) em coluna única.
- Espaçamento: cabeçalho e resumo compactos (como no Modelo);
  corpo do artigo em espaço duplo, Times New Roman 12 (regra explícita
  das Diretrizes, e também o que o estilo "Corpo do texto" do Modelo usa);
  referências em espaço simples.
- Papel A4, margens de 3 cm, páginas numeradas.
"""
import io
import re
import sys
import zipfile
import shutil
import tempfile
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RESUMO_HEADS = {"RESUMO", "ABSTRACT", "RESUMEN"}
KEYWORD_PREFIXES = ("PALAVRAS-CHAVE", "KEYWORDS", "PALABRAS-CLAVE", "PALABRAS CLAVE")
ORCID_RE = re.compile(r"(https?://orcid\.org/\S+)", re.I)
SECTION_HEAD_RE = re.compile(r"^\s*\d+(\.\d+)*\.?\s+\S")
CAPTION_RE = re.compile(r"^\s*(Quadro|Figura|Tabela|Gr[áa]fico)\s*\d+", re.I)
REFERENCES_RE = re.compile(r"^\s*(\d+(\.\d+)*\.?\s+)?(REFER[ÊE]NCIAS)\s*$", re.I)

LOGO_MEDIA = "word/media/image1.jpeg"
ICON_MEDIA = "word/media/image2.png"
BAR_MEDIA = "word/media/image3.png"


def _extract_asset(docx_path, media_path):
    try:
        with zipfile.ZipFile(docx_path) as z:
            return io.BytesIO(z.read(media_path))
    except Exception:
        return None


def _load_logo(modelo_path):
    """Prioriza um logo.png dedicado (fundo transparente, sem caixa
    branca) na mesma pasta do script; se não existir, cai de volta pro
    logo embutido no Modelo.docx."""
    here = os.path.dirname(os.path.abspath(__file__))
    local_logo = os.path.join(here, "logo.png")
    if os.path.exists(local_logo):
        with open(local_logo, "rb") as f:
            return io.BytesIO(f.read())
    return _extract_asset(modelo_path, LOGO_MEDIA) if modelo_path else None


def _extract_inline_image(paragraph):
    """Se o parágrafo contém uma imagem embutida (figura do corpo do
    artigo), devolve (bytes_da_imagem, largura_emu, altura_emu).
    Sem isso, figuras eram descartadas em silêncio — igual acontecia
    com as tabelas antes da correção."""
    for run in paragraph.runs:
        blip = run._r.find('.//' + qn('a:blip'))
        if blip is None:
            continue
        rId = blip.get(qn('r:embed'))
        if not rId:
            continue
        try:
            image_part = paragraph.part.related_parts[rId]
        except KeyError:
            continue
        extent = run._r.find('.//' + qn('wp:extent'))
        cx = int(extent.get('cx')) if extent is not None else None
        cy = int(extent.get('cy')) if extent is not None else None
        return image_part.blob, cx, cy
    return None


def classify(doc_path):
    d = Document(doc_path)
    paras = [p for p in d.paragraphs]

    blocks = {
        "titles": [], "authors": [], "abstracts": [], "body": [], "references": [],
    }

    i, n = 0, len(paras)

    while i < n and paras[i].text.strip() == "":
        i += 1
    while i < n and paras[i].text.strip() != "":
        txt = paras[i].text.strip()
        if ORCID_RE.search(txt) or txt.upper() in RESUMO_HEADS:
            break
        blocks["titles"].append(txt)
        i += 1

    while i < n:
        while i < n and paras[i].text.strip() == "":
            i += 1
        if i >= n:
            break
        txt = paras[i].text.strip()
        if txt.upper() in RESUMO_HEADS:
            break
        m = ORCID_RE.search(txt)
        if m:
            orcid_url = m.group(1).rstrip(".,;")
            name_only = txt[:m.start()].strip().rstrip("-–—:").strip()
        else:
            orcid_url, name_only = "", txt
        i += 1
        while i < n and paras[i].text.strip() == "":
            i += 1
        bio = ""
        if i < n and paras[i].text.strip().upper() not in RESUMO_HEADS:
            bio = paras[i].text.strip()
            i += 1
        blocks["authors"].append((name_only, orcid_url, bio))

    while i < n:
        while i < n and paras[i].text.strip() == "":
            i += 1
        if i >= n:
            break
        head = paras[i].text.strip().upper().rstrip(":")
        if head not in RESUMO_HEADS:
            break
        i += 1
        text_lines = []
        while i < n and paras[i].text.strip() == "":
            i += 1
        while i < n and not any(paras[i].text.strip().upper().startswith(p) for p in KEYWORD_PREFIXES):
            if paras[i].text.strip():
                text_lines.append(paras[i].text.strip())
            i += 1
        kw_label, kw_text = "", ""
        if i < n:
            kw_full = paras[i].text.strip()
            m = re.split(r":", kw_full, maxsplit=1)
            kw_label = m[0].strip()
            kw_text = m[1].strip() if len(m) > 1 else ""
            i += 1
        blocks["abstracts"].append({
            "label": head, "text": " ".join(text_lines),
            "kw_label": kw_label, "kw_text": kw_text,
        })

    # ---- A partir daqui, percorre o documento na ORDEM REAL (parágrafos
    # E tabelas), porque document.paragraphs ignora tabelas por completo —
    # era por isso que quadros/tabelas do artigo desapareciam. ----
    content = list(d.iter_inner_content())

    def is_tbl(item):
        return not hasattr(item, "text")

    # acha, na lista `content`, a posição correspondente a paras[i] (fim
    # dos resumos / início do corpo do artigo)
    body_start = len(content)
    if i < n:
        target = paras[i]._p
        for idx, item in enumerate(content):
            if not is_tbl(item) and item._p is target:
                body_start = idx
                break

    ref_idx = None
    for j in range(body_start, len(content)):
        item = content[j]
        if not is_tbl(item) and REFERENCES_RE.match(item.text.strip()):
            ref_idx = j
            break
    body_end = ref_idx if ref_idx is not None else len(content)

    for j in range(body_start, body_end):
        item = content[j]
        if is_tbl(item):
            rows = [[cell.text.strip() for cell in row.cells] for row in item.rows]
            blocks["body"].append(("table", rows, None))
            continue
        txt = item.text.strip()
        if not txt:
            img = _extract_inline_image(item)
            if img is not None:
                blocks["body"].append(("figure", img[0], (img[1], img[2])))
            continue
        is_heading = (item.style.name.startswith("Heading") or
                      bool(SECTION_HEAD_RE.match(txt) and len(txt) < 120 and txt.isupper()))
        is_caption = bool(CAPTION_RE.match(txt))
        runs_info = [(r.text, bool(r.bold), bool(r.italic)) for r in item.runs] or [(txt, False, False)]
        kind = "heading" if is_heading else ("caption" if is_caption else "para")
        blocks["body"].append((kind, txt, runs_info))

    if ref_idx is not None:
        for j in range(ref_idx + 1, len(content)):
            item = content[j]
            if is_tbl(item):
                rows = [[cell.text.strip() for cell in row.cells] for row in item.rows]
                blocks["references"].append(("table", rows))
                continue
            txt = item.text.strip()
            if txt:
                runs_info = [(r.text, bool(r.bold), bool(r.italic)) for r in item.runs] or [(txt, False, False)]
                blocks["references"].append(("para", runs_info))

    return blocks


def _set_cell_shading(cell, color):
    """Preenche o fundo de uma célula com uma cor sólida (usado pra criar
    a barrinha laranja decorativa do topo da página 1)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def _set_row_height(row, height_in, exact=True):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height_in * 1440)))
    trHeight.set(qn('w:hRule'), 'exact' if exact else 'atLeast')
    trPr.append(trHeight)


def _add_page_number_field(paragraph, font_name="Tahoma", size=11, bold=True):
    run = paragraph.add_run()
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = "PAGE"
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(instr); run._r.append(f2)


def _set_columns(sectPr, num, col1_w=None, space=None, col2_w=None):
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        docGrid = sectPr.find(qn('w:docGrid'))
        if docGrid is not None:
            docGrid.addprevious(cols)
        else:
            sectPr.append(cols)
    for child in list(cols):
        cols.remove(child)
    for attr in list(cols.attrib):
        del cols.attrib[attr]
    cols.set(qn('w:num'), str(num))
    if num == 2 and col1_w:
        cols.set(qn('w:space'), str(space or 400))
        cols.set(qn('w:equalWidth'), '0')
        c1 = OxmlElement('w:col')
        c1.set(qn('w:w'), str(col1_w))
        c1.set(qn('w:space'), str(space or 400))
        cols.append(c1)
        if col2_w:
            c2 = OxmlElement('w:col')
            c2.set(qn('w:w'), str(col2_w))
            cols.append(c2)


def _add_column_break(document):
    p = document.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'column')
    r._r.append(br)
    return p


def _hyperlink_last_picture(run, url):
    part = run.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    drawing = run._r.find(qn('w:drawing'))
    if drawing is None:
        return
    ns_wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    ns_pic = "http://schemas.openxmlformats.org/drawingml/2006/picture"
    docPr = drawing.find(f'.//{{{ns_wp}}}docPr')
    cNvPr = drawing.find(f'.//{{{ns_pic}}}cNvPr')
    for el in (docPr, cNvPr):
        if el is None:
            continue
        hlink = OxmlElement('a:hlinkClick')
        hlink.set(qn('r:id'), r_id)
        el.append(hlink)


_FOOTNOTE_MARK_PREFIX = "@@FOOTNOTE_"


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Zera as margens internas da célula (por padrão o Word reserva um
    respiro interno que, somado à largura exata da imagem, cortava a
    logo)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement(f'w:{tag}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_table_indent_zero(table):
    """Garante que a tabela comece exatamente na margem esquerda, sem
    nenhum recuo padrão."""
    tblPr = table._tbl.tblPr
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), '0')
    tblInd.set(qn('w:type'), 'dxa')
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblW.addnext(tblInd)
    else:
        tblPr.insert(0, tblInd)


def _set_cell_border(cell, edge, color="000000", size=12):
    """Adiciona uma borda colorida a um lado da célula (usado para a
    linha laranja que separa a logo do título, igual ao Modelo)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    tag = f'w:{edge}'
    border_el = tcBorders.find(qn(tag))
    if border_el is None:
        border_el = OxmlElement(tag)
        tcBorders.append(border_el)
    border_el.set(qn('w:val'), 'single')
    border_el.set(qn('w:sz'), str(size))
    border_el.set(qn('w:space'), '4')
    border_el.set(qn('w:color'), color)


def _add_footnote_marker(paragraph, index):
    r = paragraph.add_run(f"{_FOOTNOTE_MARK_PREFIX}{index}@@")
    r.font.superscript = True
    r.font.name = "Tahoma"
    r.font.size = Pt(10)
    return r


def _inject_footnotes(docx_path, footnote_texts):
    if not footnote_texts:
        return

    tmp_dir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(docx_path) as z:
            z.extractall(tmp_dir)

        NS = ('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
              'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"')
        parts = [f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:footnotes {NS}>']
        parts.append(
            '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr>'
            '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
            '<w:r><w:separator/></w:r></w:p></w:footnote>'
        )
        parts.append(
            '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr>'
            '<w:spacing w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>'
            '<w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
        )
        for idx, text in enumerate(footnote_texts, start=1):
            safe = (text.replace("&", "&amp;").replace("<", "&lt;")
                        .replace(">", "&gt;"))
            parts.append(
                f'<w:footnote w:id="{idx}"><w:p><w:pPr>'
                f'<w:spacing w:after="0" w:line="240" w:lineRule="auto"/>'
                f'<w:jc w:val="both"/></w:pPr>'
                f'<w:r><w:rPr><w:vertAlign w:val="superscript"/>'
                f'<w:rFonts w:ascii="Tahoma" w:hAnsi="Tahoma"/>'
                f'<w:sz w:val="20"/></w:rPr><w:footnoteRef/></w:r>'
                f'<w:r><w:rPr><w:rFonts w:ascii="Tahoma" w:hAnsi="Tahoma"/>'
                f'<w:sz w:val="20"/></w:rPr><w:t xml:space="preserve"> {safe}</w:t></w:r>'
                f'</w:p></w:footnote>'
            )
        parts.append('</w:footnotes>')
        with open(os.path.join(tmp_dir, 'word', 'footnotes.xml'), 'w', encoding='utf-8') as f:
            f.write(''.join(parts))

        ct_path = os.path.join(tmp_dir, '[Content_Types].xml')
        with open(ct_path, encoding='utf-8') as f:
            ct = f.read()
        if 'footnotes.xml' not in ct:
            override = ('<Override PartName="/word/footnotes.xml" '
                         'ContentType="application/vnd.openxmlformats-officedocument'
                         '.wordprocessingml.footnotes+xml"/>')
            ct = ct.replace('</Types>', override + '</Types>')
            with open(ct_path, 'w', encoding='utf-8') as f:
                f.write(ct)

        rels_path = os.path.join(tmp_dir, 'word', '_rels', 'document.xml.rels')
        with open(rels_path, encoding='utf-8') as f:
            rels = f.read()
        existing_ids = re.findall(r'Id="rId(\d+)"', rels)
        next_id = max(int(x) for x in existing_ids) + 1 if existing_ids else 1
        rel = (f'<Relationship Id="rId{next_id}" '
               f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" '
               f'Target="footnotes.xml"/>')
        rels = rels.replace('</Relationships>', rel + '</Relationships>')
        with open(rels_path, 'w', encoding='utf-8') as f:
            f.write(rels)

        doc_path = os.path.join(tmp_dir, 'word', 'document.xml')
        with open(doc_path, encoding='utf-8') as f:
            doc_xml = f.read()
        for idx in range(1, len(footnote_texts) + 1):
            marker = f"{_FOOTNOTE_MARK_PREFIX}{idx}@@"
            old_t = f'<w:t>{marker}</w:t>'
            new_el = f'<w:footnoteReference w:id="{idx}"/>'
            doc_xml = doc_xml.replace(old_t, new_el)
        with open(doc_path, 'w', encoding='utf-8') as f:
            f.write(doc_xml)

        out_tmp = docx_path + '.tmp'
        if os.path.exists(out_tmp):
            os.remove(out_tmp)
        with zipfile.ZipFile(out_tmp, 'w', zipfile.ZIP_DEFLATED) as zf:
            for root, _, files in os.walk(tmp_dir):
                for fname in files:
                    full = os.path.join(root, fname)
                    arc = os.path.relpath(full, tmp_dir)
                    zf.write(full, arc)
        shutil.move(out_tmp, docx_path)
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _add_floating_image(paragraph, image_stream, width, height):
    """Insere a logo como imagem flutuante, fixa no canto superior
    esquerdo, sem contornar texto (o alinhamento do título ao lado dela
    é feito por recuo de parágrafo, não pelo wrap automático do Word —
    isso evita que a última linha "escape" do alinhamento em títulos
    longos)."""
    run = paragraph.add_run()
    run.add_picture(image_stream, width=width, height=height)
    drawing = run._r.find(qn('w:drawing'))
    inline = drawing.find(qn('wp:inline'))
    extent = inline.find(qn('wp:extent'))
    docPr = inline.find(qn('wp:docPr'))
    graphic = inline.find(qn('a:graphic'))

    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', '0'); anchor.set('distB', '0')
    anchor.set('distL', '0'); anchor.set('distR', '0')
    anchor.set('simplePos', '0'); anchor.set('relativeHeight', '1')
    anchor.set('behindDoc', '0'); anchor.set('locked', '0')
    anchor.set('layoutInCell', '1'); anchor.set('allowOverlap', '1')

    simplePos = OxmlElement('wp:simplePos'); simplePos.set('x', '0'); simplePos.set('y', '0')
    posH = OxmlElement('wp:positionH'); posH.set('relativeFrom', 'column')
    offH = OxmlElement('wp:posOffset'); offH.text = '0'
    posH.append(offH)
    posV = OxmlElement('wp:positionV'); posV.set('relativeFrom', 'paragraph')
    offV = OxmlElement('wp:posOffset'); offV.text = '0'
    posV.append(offV)
    wrap = OxmlElement('wp:wrapNone')

    anchor.append(simplePos)
    anchor.append(posH)
    anchor.append(posV)
    anchor.append(extent)
    effectExtent = OxmlElement('wp:effectExtent')
    for a in ('l', 't', 'r', 'b'):
        effectExtent.set(a, '0')
    anchor.append(effectExtent)
    anchor.append(wrap)
    anchor.append(docPr)
    cNvGraphicFramePr = OxmlElement('wp:cNvGraphicFramePr')
    anchor.append(cNvGraphicFramePr)
    anchor.append(graphic)

    drawing.remove(inline)
    drawing.append(anchor)
    return run


def _add_body_table(document, rows_data, usable_in):
    """Recria um quadro/tabela do artigo original (ex.: 'Quadro 1') como
    uma tabela de verdade no documento final, com bordas e cabeçalho em
    negrito."""
    if not rows_data or not rows_data[0]:
        return
    n_rows = len(rows_data)
    n_cols = len(rows_data[0])
    table = document.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    table.autofit = True
    col_w = Inches(usable_in / n_cols)
    for col in table.columns:
        col.width = col_w
    for ridx, row_vals in enumerate(rows_data):
        row = table.rows[ridx]
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)
        for cidx, val in enumerate(row_vals):
            if cidx >= n_cols:
                continue
            cell = table.cell(ridx, cidx)
            cell.width = col_w
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(val)
            r.font.name = "Tahoma"
            r.font.size = Pt(9)
            r.bold = (ridx == 0)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)


def build_document(blocks, out_path, modelo_path=None):
    d = Document()

    logo_stream = _load_logo(modelo_path)
    icon_stream = _extract_asset(modelo_path, ICON_MEDIA) if modelo_path else None

    sec = d.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.01)
    sec.right_margin = Cm(1.78)
    sec.top_margin = Cm(1.69)
    sec.bottom_margin = Cm(2.54)
    sec.header_distance = Cm(0)
    sec.footer_distance = Cm(1.27)

    normal = d.styles["Normal"]
    normal.font.name = "Tahoma"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(0)

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_page_number_field(fp)

    usable_w_in = (21 - 2.01 - 1.78) / 2.54

    # ---- Cabeçalho corrido (páginas 2+): barrinha laranja (imagem real)
    # acima de "Revista Educar Mais", com "| 2026 | Volume 10 |" embaixo
    # dela (à esquerda), e "CC BY-NC 4.0" / "e-ISSN 2237-9185" alinhados
    # na mesma linha do nome da revista. A página 1 fica sem esse
    # cabeçalho, porque já tem a logo/título ocupando esse espaço. ----
    sec.different_first_page_header_footer = True
    header = sec.header
    header.is_linked_to_previous = False
    bar_stream = _extract_asset(modelo_path, BAR_MEDIA) if modelo_path else None

    bar_p = header.paragraphs[0]
    # o Word sempre cria um parágrafo vazio ao criar o cabeçalho; ele ficava
    # empurrando a barra pra baixo. Usa esse mesmo parágrafo pra barra
    # (em vez de criar um novo), evitando o parágrafo fantasma.
    if bar_stream:
        bar_stream.seek(0)
        bar_p.paragraph_format.space_after = Pt(2)
        bar_p.add_run().add_picture(bar_stream, width=Inches(1.68), height=Inches(0.2))

    hp1 = header.add_paragraph()
    hp1.paragraph_format.space_before = Pt(2)
    hp1.paragraph_format.space_after = Pt(0)
    for txt in (" Revista Educar Mais", " " * 49 + "CC BY-NC 4.0", " " * 35 + "e-ISSN 2237-9185"):
        r = hp1.add_run(txt)
        r.font.name = "Tahoma"
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0x99, 0x00)

    hp2 = header.add_paragraph()
    hp2.paragraph_format.space_before = Pt(0)
    hp2.paragraph_format.space_after = Pt(6)
    r2 = hp2.add_run(" " * 11 + "| 2026 | Volume 10 |")
    r2.font.name = "Tahoma"
    r2.font.size = Pt(6)
    r2.font.color.rgb = RGBColor(0xFF, 0x99, 0x00)

    # ---- Topo da página 1: barrinha laranja (imagem de verdade, extraída
    # do Modelo) + "2026 | Volume 10 | Pág. 1 a XX" + tab + espaços +
    # "DOI: ..." — tudo em um parágrafo único, igual à estrutura real do
    # Modelo (não uma tabela de 2 colunas). ----
    fp_header = sec.first_page_header
    bar_p = fp_header.paragraphs[0]
    bar_p.paragraph_format.space_after = Pt(2)
    bar_p.paragraph_format.space_before = Pt(0)
    bar_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    bar_p.paragraph_format.line_spacing = Pt(14.2)
    if bar_stream:
        bar_stream.seek(0)
        bar_run = bar_p.add_run()
        bar_run.add_picture(bar_stream, width=Cm(4.28), height=Cm(0.5))

    line_p = fp_header.add_paragraph()
    line_p.paragraph_format.space_before = Pt(2)
    line_p.paragraph_format.space_after = Pt(0)
    vr = line_p.add_run("           2026 | Volume 10 | Pág. 1 a XX")
    vr.font.name = "Tahoma"
    vr.font.size = Pt(6)
    vr.font.color.rgb = RGBColor(0xED, 0x7D, 0x31)
    tab_r = line_p.add_run()
    tab_el = OxmlElement('w:tab')
    tab_r._r.append(tab_el)
    doi_r = line_p.add_run(" " * 43 + "DOI: https://doi.org/10.15536/reducarmais.10.2026.xxxxx")
    doi_r.font.name = "Tahoma"
    doi_r.font.size = Pt(8)
    doi_r.bold = True
    doi_r.font.color.rgb = RGBColor(0xED, 0x7D, 0x31)

    usable_w = 21 - 2.01 - 1.78

    def para(text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             size=12, space_after=4, line_spacing=1.15):
        p = d.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = line_spacing
        if text:
            r = p.add_run(text)
            r.font.name = "Tahoma"
            r.font.size = Pt(size)
            r.bold = bold
            r.italic = italic
        return p

    # ---- Logo + título em tabela sem bordas (1 linha, 2 colunas).
    # Muito mais robusto que imagem flutuante: a célula do título nunca
    # perde o alinhamento, e o que vem depois (autores) sempre começa
    # abaixo da tabela inteira, sem sobrepor a logo. ----
    LOGO_WIDTH_IN = 1.3
    usable_in = usable_w / 2.54
    title_col_in = usable_in - LOGO_WIDTH_IN - 0.15

    titles = blocks["titles"]
    table = d.add_table(rows=1, cols=2)
    table.autofit = False
    _set_table_indent_zero(table)
    logo_cell, title_cell = table.rows[0].cells
    logo_cell.width = Inches(LOGO_WIDTH_IN)
    title_cell.width = Inches(title_col_in)
    for col, w in zip(table.columns, (Inches(LOGO_WIDTH_IN), Inches(title_col_in))):
        col.width = w
    _set_cell_border(title_cell, 'left', color="FF7F26", size=18)
    _set_cell_margins(logo_cell, top=0, bottom=0, left=0, right=0)
    TITLE_CELL_LEFT_MARGIN_IN = 0.15
    _set_cell_margins(title_cell, top=0, bottom=0,
                       left=int(TITLE_CELL_LEFT_MARGIN_IN * 1440), right=100)

    logo_p = logo_cell.paragraphs[0]
    logo_p.paragraph_format.space_after = Pt(0)
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_stream:
        logo_p.add_run().add_picture(logo_stream, width=Inches(LOGO_WIDTH_IN - 0.1))

    if titles:
        tp = title_cell.paragraphs[0]
        tp.paragraph_format.space_after = Pt(6)
        tp.paragraph_format.line_spacing = 1.15
        r = tp.add_run(titles[0])
        r.font.name = "Tahoma"
        r.font.size = Pt(15)
        r.bold = True
    for t in titles[1:]:
        tp = title_cell.add_paragraph()
        tp.paragraph_format.space_after = Pt(6)
        tp.paragraph_format.line_spacing = 1.15
        r = tp.add_run(t)
        r.font.name = "Tahoma"
        r.font.size = Pt(12)
        r.bold = True
        r.italic = True

    footnote_texts = []
    authors_p = d.add_paragraph()
    authors_p.paragraph_format.left_indent = Inches(LOGO_WIDTH_IN + TITLE_CELL_LEFT_MARGIN_IN)
    authors_p.paragraph_format.space_before = Pt(10)
    authors_p.paragraph_format.space_after = Pt(10)
    authors_p.paragraph_format.line_spacing = 1.2
    for idx, (name_only, orcid_url, bio) in enumerate(blocks["authors"]):
        r = authors_p.add_run(name_only)
        r.font.name = "Tahoma"
        r.font.size = Pt(11)
        if bio:
            footnote_texts.append(bio)
            _add_footnote_marker(authors_p, len(footnote_texts))
        if orcid_url and icon_stream:
            icon_stream.seek(0)
            icon_run = authors_p.add_run()
            icon_run.add_picture(icon_stream, width=Inches(0.26))
            _hyperlink_last_picture(icon_run, orcid_url)
        if idx < len(blocks["authors"]) - 1:
            sep = authors_p.add_run("  ꞏ  ")
            sep.font.name = "Tahoma"
            sep.font.size = Pt(11)

    for ab in blocks["abstracts"]:
        is_pt = ab["label"] == "RESUMO"
        para(ab["label"], bold=True, italic=not is_pt, align=WD_ALIGN_PARAGRAPH.LEFT,
             size=10, space_after=4, line_spacing=1.15)
        body_p = d.add_paragraph()
        body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_p.paragraph_format.line_spacing = 1.15
        body_p.paragraph_format.space_after = Pt(6)
        r = body_p.add_run(ab["text"])
        r.font.name = "Tahoma"
        r.font.size = Pt(10)
        r.italic = not is_pt
        if ab["kw_label"]:
            kw_p = d.add_paragraph()
            kw_p.paragraph_format.line_spacing = 1.15
            kw_p.paragraph_format.space_after = Pt(12)
            r1 = kw_p.add_run(ab["kw_label"] + ": ")
            r1.bold = True
            r1.italic = not is_pt
            r1.font.name = "Tahoma"
            r1.font.size = Pt(10)
            r2 = kw_p.add_run(ab["kw_text"])
            r2.italic = not is_pt
            r2.font.name = "Tahoma"
            r2.font.size = Pt(10)

    for item in blocks["body"]:
        kind = item[0]
        if kind == "table":
            _add_body_table(d, item[1], usable_in)
            continue
        if kind == "figure":
            img_bytes, (cx, cy) = item[1], item[2]
            fig_p = d.add_paragraph()
            fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fig_p.paragraph_format.space_before = Pt(6)
            fig_p.paragraph_format.space_after = Pt(8)
            fig_p.paragraph_format.keep_with_next = True
            max_w = Inches(usable_in)
            if cx:
                width = min(Inches(cx / 914400), max_w)
            else:
                width = max_w
            fig_p.add_run().add_picture(io.BytesIO(img_bytes), width=width)
            continue
        if kind == "heading":
            txt = item[1]
            p = d.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.keep_with_next = True
            r = p.add_run(txt.upper())
            r.bold = True
            r.font.name = "Tahoma"
            r.font.size = Pt(11)
        elif kind == "caption":
            txt, runs_info = item[1], item[2]
            p = d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.1
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            for rtext, rbold, ritalic in runs_info:
                if not rtext:
                    continue
                r = p.add_run(rtext)
                r.font.name = "Tahoma"
                r.font.size = Pt(10)
                r.bold = rbold
                r.italic = ritalic
        else:
            txt, runs_info = item[1], item[2]
            p = d.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(8)
            for rtext, rbold, ritalic in runs_info:
                if not rtext:
                    continue
                r = p.add_run(rtext)
                r.font.name = "Tahoma"
                r.font.size = Pt(11)
                r.bold = rbold
                r.italic = ritalic

    if blocks["references"]:
        hp = d.add_paragraph()
        hp.paragraph_format.space_before = Pt(18)
        hp.paragraph_format.space_after = Pt(6)
        hp.paragraph_format.line_spacing = 1.0
        hp.paragraph_format.keep_with_next = True
        rh = hp.add_run("REFERÊNCIAS")
        rh.bold = True
        rh.font.name = "Tahoma"
        rh.font.size = Pt(11)

        for ref_item in blocks["references"]:
            if ref_item[0] == "table":
                _add_body_table(d, ref_item[1], usable_in)
                continue
            ref_runs = ref_item[1]
            p = d.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(6)
            for rtext, rbold, ritalic in ref_runs:
                if not rtext:
                    continue
                r = p.add_run(rtext)
                r.font.name = "Tahoma"
                r.font.size = Pt(11)
                r.bold = rbold
                r.italic = ritalic

    d.save(out_path)
    _inject_footnotes(out_path, footnote_texts)


def convert(in_path, out_path, modelo_path=None):
    blocks = classify(in_path)
    build_document(blocks, out_path, modelo_path=modelo_path)
    return blocks


if __name__ == "__main__":
    modelo = sys.argv[3] if len(sys.argv) > 3 else None
    convert(sys.argv[1], sys.argv[2], modelo_path=modelo)
